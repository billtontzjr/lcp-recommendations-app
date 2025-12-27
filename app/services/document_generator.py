"""Word document generation service - matching Farress app formatting."""
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from datetime import datetime
from app.utils.currency import format_currency

# Constants
GRAY_BACKGROUND = "D3D3D3"
FONT_NAME = "Times New Roman"


def set_cell_background(cell, color):
    """Set the background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def set_cell_padding(cell):
    """Set padding (margins) for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ['top', 'bottom']:
        margin = OxmlElement(f'w:{side}')
        margin.set(qn('w:w'), "70")
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)


def set_padding_for_table(table):
    """Set padding for all cells in a table."""
    for row in table.rows:
        for cell in row.cells:
            set_cell_padding(cell)


def set_bold_borders(table):
    """Set bold borders for a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)

    border_styles = {
        "top": "single",
        "left": "single",
        "bottom": "single",
        "right": "single",
        "insideH": "single",
        "insideV": "single"
    }

    for border_name, border_type in border_styles.items():
        border = tblBorders.find(qn(f'w:{border_name}'))
        if border is None:
            border = OxmlElement(f'w:{border_name}')
            tblBorders.append(border)
        border.set(qn('w:val'), border_type)
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')


def add_empty_gray_row(table):
    """Add an empty gray row to the table with only left/right bold borders."""
    empty_row = table.add_row()

    for cell in empty_row.cells:
        set_cell_background(cell, GRAY_BACKGROUND)
        cell.text = ""
        # Set minimal height for the gray row
        set_cell_padding(cell)

    tbl = table._tbl
    last_row = tbl.findall(qn("w:tr"))[-1]
    cells = last_row.findall(qn("w:tc"))

    for idx, tc in enumerate(cells):
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)

        # Remove existing borders element if present
        existing_borders = tcPr.find(qn("w:tcBorders"))
        if existing_borders is not None:
            tcPr.remove(existing_borders)

        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

        # Set all borders to nil first
        for border_name in ["top", "bottom", "left", "right"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "nil")
            tcBorders.append(border)

        # Bold left border for first cell only
        if idx == 0:
            left_border = tcBorders.find(qn("w:left"))
            if left_border is not None:
                tcBorders.remove(left_border)
            left_border = OxmlElement("w:left")
            left_border.set(qn("w:val"), "single")
            left_border.set(qn("w:sz"), "12")
            left_border.set(qn("w:color"), "000000")
            tcBorders.append(left_border)

        # Bold right border for last cell only
        if idx == len(cells) - 1:
            right_border = tcBorders.find(qn("w:right"))
            if right_border is not None:
                tcBorders.remove(right_border)
            right_border = OxmlElement("w:right")
            right_border.set(qn("w:val"), "single")
            right_border.set(qn("w:sz"), "12")
            right_border.set(qn("w:color"), "000000")
            tcBorders.append(right_border)


def format_cell_text(cell, text, font_size=Pt(10), bold=False, center=False, gray=False):
    """Format a cell with text and styling."""
    cell.text = str(text) if text is not None else ''

    if cell.paragraphs and cell.paragraphs[0].runs:
        run = cell.paragraphs[0].runs[0]
        run.font.name = FONT_NAME
        run.font.size = font_size
        run.font.bold = bold

    if center:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if gray:
        set_cell_background(cell, GRAY_BACKGROUND)

    set_cell_padding(cell)


def format_date(date_val):
    """Format date for display."""
    if not date_val:
        return ''
    if isinstance(date_val, datetime):
        return date_val.strftime('%B %d, %Y')
    if isinstance(date_val, str):
        return date_val
    return str(date_val)


def format_cost(value):
    """Format cost value with commas and 2 decimal places."""
    if value is None:
        return "$0.00"
    try:
        val = float(value)
        return f"${val:,.2f}"
    except (ValueError, TypeError):
        return "$0.00"


def add_new_page(doc):
    """Add a page break."""
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def generate_lcp_document(patient_info, cost_data, output_path):
    """
    Generate the Life Care Plan Recommendations Word document.
    """
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(10)

    # Set margins
    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

    # Page 1: Title Page
    add_title_page(doc, patient_info)

    # Page 2: Appendix A Summary
    add_new_page(doc)
    add_appendix_a(doc, patient_info, cost_data)

    # Section Pages (one per category)
    for category, data in cost_data['category_totals'].items():
        add_new_page(doc)
        add_section_page(doc, patient_info, category, data, cost_data['totals'])

    doc.save(output_path)
    return output_path


def add_title_page(doc, patient_info):
    """Add the title page."""
    # Main title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('LIFE CARE PLAN RECOMMENDATIONS')
    run.font.name = FONT_NAME
    run.font.size = Pt(16)
    run.font.bold = True

    doc.add_paragraph()

    # Patient name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run(patient_info.get('patient_name', ''))
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_paragraph()

    # Patient details
    details = [
        ('Date of Birth:', format_date(patient_info.get('date_of_birth'))),
        ('Date of Injury:', format_date(patient_info.get('date_of_injury'))),
        ('Date of Report:', format_date(patient_info.get('date_of_report'))),
        ('Life Expectancy:', f"{patient_info.get('life_expectancy', '')} years"),
    ]

    for label, value in details:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f'{label} ')
        run.font.name = FONT_NAME
        run.font.size = Pt(11)
        run = para.add_run(str(value))
        run.font.name = FONT_NAME
        run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    # Prepared by
    prep = doc.add_paragraph()
    prep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prep.add_run('Prepared by:')
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.bold = True

    doc.add_paragraph()

    credentials = [
        'William Tontz, Jr., MD',
        'Board-Certified Orthopedic Surgeon',
        'Certified Life Care Planner (CLCP)',
        '',
        'Precision Life Care Planning',
    ]

    for line in credentials:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        run.font.name = FONT_NAME
        run.font.size = Pt(11)


def add_appendix_a(doc, patient_info, cost_data):
    """Add Appendix A summary page."""
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run('APPENDIX A')
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_paragraph()

    # Client info with tabs
    client_info = [
        ("Client:", patient_info.get('patient_name', '')),
        ("Date of Birth:", format_date(patient_info.get('date_of_birth'))),
        ("Date of Injury:", format_date(patient_info.get('date_of_injury'))),
        ("Date of Report:", format_date(patient_info.get('date_of_report'))),
        ("Life Expectancy:", f"{patient_info.get('life_expectancy', '')} years"),
    ]

    for title, value in client_info:
        para = doc.add_paragraph()
        para.paragraph_format.tab_stops.add_tab_stop(Inches(1.5))
        run = para.add_run(title)
        run.font.name = FONT_NAME
        run.font.bold = True
        para.add_run("\t")
        run2 = para.add_run(str(value))
        run2.font.name = FONT_NAME
        para.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # Main summary table
    category_totals = cost_data['category_totals']
    totals = cost_data['totals']

    # Filter categories with non-zero costs
    filtered_categories = {
        cat: data for cat, data in category_totals.items()
        if data['annual_cost'] != 0 or data['one_time_cost'] != 0
    }

    # Create main table
    num_data_rows = len(filtered_categories)
    table = doc.add_table(rows=2 + num_data_rows, cols=5)
    table.style = 'Table Grid'

    # Column widths
    col_widths = [Inches(1.5), Inches(1.2), Inches(1.5), Inches(2.5), Inches(1.5)]
    for row in table.rows:
        for i, width in enumerate(col_widths):
            row.cells[i].width = width

    # Row 0: Title (merged)
    title_cell = table.rows[0].cells[0]
    for i in range(1, 5):
        title_cell.merge(table.rows[0].cells[i])
    format_cell_text(title_cell, 'Lifetime Projected Costs', Pt(11), bold=True, center=True, gray=True)

    # Row 1: Headers
    headers = ['Section', 'Annual Cost', 'Multiplied Annual Cost', 'One Time Cost', 'Total Lifetime Cost']
    for i, header_text in enumerate(headers):
        format_cell_text(table.rows[1].cells[i], header_text, Pt(10), bold=True, center=True, gray=True)

    # Data rows
    row_idx = 2
    life_exp = float(totals.get('life_expectancy', 0) or 0)

    for category, data in filtered_categories.items():
        annual = data['annual_cost']
        one_time = data['one_time_cost']
        multiplied = annual * life_exp
        total_lifetime = multiplied + one_time

        format_cell_text(table.rows[row_idx].cells[0], category, Pt(10))
        format_cell_text(table.rows[row_idx].cells[1], f"{annual:,.2f}", Pt(10), center=True)
        format_cell_text(table.rows[row_idx].cells[2], f"{multiplied:,.2f}", Pt(10), center=True)
        format_cell_text(table.rows[row_idx].cells[3], f"{one_time:,.2f}", Pt(10), center=True)
        format_cell_text(table.rows[row_idx].cells[4], f"{total_lifetime:,.2f}", Pt(10), center=True)
        row_idx += 1

    # Add empty gray row
    add_empty_gray_row(table)
    set_bold_borders(table)

    # Totals table (3 columns) - no paragraph gap
    totals_table = doc.add_table(rows=5, cols=3)
    totals_table.style = 'Table Grid'

    # Totals column widths
    totals_col_widths = [Inches(1), Inches(4.2), Inches(2.0)]
    for row in totals_table.rows:
        for i, width in enumerate(totals_col_widths):
            row.cells[i].width = width

    totals_data = [
        ('', 'Total Annual Cost:', f"{totals['total_annual']:,.2f}"),
        ('', f"Lifetime Annual Cost (Annual x {totals['life_expectancy']} years):", f"{totals['lifetime_annual']:,.2f}"),
        ('', 'Total Lifetime One Time Cost:', f"{totals['total_one_time']:,.2f}"),
        ('', '', ''),  # Empty row
        ('', 'Grand Total (total lifetime cost):', f"{totals['grand_total']:,.2f}"),
    ]

    for i, (col1, col2, col3) in enumerate(totals_data):
        row = totals_table.rows[i]
        format_cell_text(row.cells[0], col1, Pt(10), bold=True)

        if i == 3:  # Empty row
            for cell in row.cells:
                cell.text = ''
                set_cell_padding(cell)
        else:
            is_grand = (i == 4)
            format_cell_text(row.cells[1], col2, Pt(8) if not is_grand else Pt(10), bold=is_grand)
            format_cell_text(row.cells[2], col3, Pt(10), bold=True, center=True)

    set_bold_borders(totals_table)
    set_padding_for_table(totals_table)


def add_section_page(doc, patient_info, category, category_data, totals):
    """Add a section page with data table, rationale, and sources."""
    # Client info header (small)
    client_info = [
        ("Client:", patient_info.get('patient_name', '')),
        ("Date of Birth:", format_date(patient_info.get('date_of_birth'))),
        ("Date of Injury:", format_date(patient_info.get('date_of_injury'))),
        ("Date of Report:", format_date(patient_info.get('date_of_report'))),
    ]

    for title, value in client_info:
        para = doc.add_paragraph()
        para.paragraph_format.tab_stops.add_tab_stop(Inches(1.5))
        run = para.add_run(title)
        run.font.name = FONT_NAME
        run.font.bold = True
        para.add_run("\t")
        run2 = para.add_run(str(value))
        run2.font.name = FONT_NAME
        para.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    items = category_data['items']
    age_init = patient_info.get('age_initiated', '')
    until_age = patient_info.get('until_age', '')

    # Section data table (7 columns)
    num_rows = 2 + len(items)  # title + headers + data rows
    table = doc.add_table(rows=num_rows, cols=7)
    table.style = 'Table Grid'

    # Column widths
    col_widths = [Inches(2), Inches(1), Inches(1), Inches(1), Inches(1.5), Inches(1.5), Inches(1.5)]
    for row in table.rows:
        for i, width in enumerate(col_widths):
            row.cells[i].width = width

    # Row 0: Category title (merged)
    title_cell = table.rows[0].cells[0]
    for i in range(1, 7):
        title_cell.merge(table.rows[0].cells[i])
    format_cell_text(title_cell, category, Pt(11), bold=True, center=True, gray=True)

    # Row 1: Headers
    headers = ['Item', 'Age Initiated', 'Until Age', 'Cost', 'Frequency of Visit', 'Annual Cost', 'One Time Cost']
    for i, header_text in enumerate(headers):
        format_cell_text(table.rows[1].cells[i], header_text, Pt(10), bold=True, center=True, gray=True)

    # Data rows
    for row_idx, item in enumerate(items, start=2):
        item_name = item.get('item', '') or item.get('service_description', '')
        format_cell_text(table.rows[row_idx].cells[0], item_name, Pt(10))
        format_cell_text(table.rows[row_idx].cells[1], age_init, Pt(10), center=True)
        format_cell_text(table.rows[row_idx].cells[2], until_age, Pt(10), center=True)
        format_cell_text(table.rows[row_idx].cells[3], f"{item.get('unit_cost', 0):,.2f}", Pt(10), center=True)
        format_cell_text(table.rows[row_idx].cells[4], item.get('frequency', ''), Pt(10), center=True)
        format_cell_text(table.rows[row_idx].cells[5], f"{item.get('annual_cost', 0):,.2f}", Pt(10), center=True)
        format_cell_text(table.rows[row_idx].cells[6], f"{item.get('one_time_cost', 0):,.2f}", Pt(10), center=True)

    # Add empty gray row
    add_empty_gray_row(table)
    set_bold_borders(table)

    # Totals row table (2 columns) - no paragraph gap
    section_annual = category_data['annual_cost']
    section_onetime = category_data['one_time_cost']

    totals_section = doc.add_table(rows=2, cols=2)
    totals_section.style = 'Table Grid'

    for row in totals_section.rows:
        row.cells[0].width = Inches(4.0)
        row.cells[1].width = Inches(4.0)

    format_cell_text(totals_section.rows[0].cells[0], 'Total Annual Cost:', Pt(10), bold=True, center=True)
    format_cell_text(totals_section.rows[0].cells[1], f"{section_annual:,.2f}", Pt(10), bold=True, center=True)
    format_cell_text(totals_section.rows[1].cells[0], 'Total One Time Cost:', Pt(10), bold=True, center=True)
    format_cell_text(totals_section.rows[1].cells[1], f"{section_onetime:,.2f}", Pt(10), bold=True, center=True)

    set_bold_borders(totals_section)

    # Rationale table - no paragraph gap
    rationales = [item.get('rationale', '') for item in items if item.get('rationale')]
    rationale_text = ' '.join(rationales) if rationales else ''

    rationale_table = doc.add_table(rows=2, cols=1)
    rationale_table.style = 'Table Grid'
    rationale_table.rows[0].cells[0].width = Inches(8.0)
    rationale_table.rows[1].cells[0].width = Inches(8.0)

    format_cell_text(rationale_table.rows[0].cells[0], 'Rationale', Pt(10), bold=True, gray=True)
    format_cell_text(rationale_table.rows[1].cells[0], rationale_text, Pt(10))

    set_bold_borders(rationale_table)

    # Sources table - no paragraph gap
    sources = [item.get('source', '') for item in items if item.get('source')]
    sources_text = '; '.join(set(sources)) if sources else ''

    sources_table = doc.add_table(rows=2, cols=1)
    sources_table.style = 'Table Grid'
    sources_table.rows[0].cells[0].width = Inches(8.0)
    sources_table.rows[1].cells[0].width = Inches(8.0)

    format_cell_text(sources_table.rows[0].cells[0], 'Sources', Pt(10), bold=True, gray=True)
    format_cell_text(sources_table.rows[1].cells[0], sources_text, Pt(10))

    set_bold_borders(sources_table)
